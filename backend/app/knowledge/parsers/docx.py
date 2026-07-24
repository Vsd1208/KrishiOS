"""DOCX parser using python-docx.

python-docx reads OOXML format (.docx) files and provides paragraph-level
access to text content. Tables are also traversed to ensure no content is lost.

Since DOCX files do not have fixed page boundaries, the entire document
is treated as a single page (page_number=1). This matches the expected
interface: downstream chunking handles segmentation regardless of source.

Design note: .doc (legacy binary) format is NOT supported. Users must
convert .doc → .docx before upload. This is enforced at the MIME type
validation layer in the upload endpoint.
"""

from __future__ import annotations

import io
import time

from loguru import logger

from app.knowledge.interfaces.parser import BaseParser, ParsedDocument, ParsedPage

try:
    import docx  # python-docx
except ImportError as exc:
    msg = "python-docx is required for DOCX parsing. Install with: pip install python-docx"
    raise ImportError(msg) from exc


class DocxParser:
    """DOCX text extractor using python-docx.

    Implements the BaseParser protocol.
    Extracts paragraphs and table cell content in document order.
    """

    _SUPPORTED_MIMES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/docx",
    }

    def supports(self, mime_type: str) -> bool:
        return mime_type.lower() in self._SUPPORTED_MIMES

    async def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Extract text from a DOCX file.

        Parameters
        ----------
        file_bytes:
            Raw DOCX binary content.
        filename:
            Original filename (used for logging only).

        Returns
        -------
        ParsedDocument
            A single ParsedPage containing all document text.
        """
        t0 = time.perf_counter()

        try:
            document = docx.Document(io.BytesIO(file_bytes))
        except Exception as exc:
            logger.error("DocxParser: failed to open '{}': {}", filename, exc)
            raise

        text_parts: list[str] = []

        # ── Paragraphs ────────────────────────────────────────────────────
        for para in document.paragraphs:
            stripped = para.text.strip()
            if stripped:
                text_parts.append(stripped)

        # ── Tables ────────────────────────────────────────────────────────
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        elapsed = time.perf_counter() - t0

        logger.info(
            "DocxParser: '{}' → {} chars in {:.3f}s",
            filename,
            len(full_text),
            elapsed,
        )

        return ParsedDocument(
            pages=[ParsedPage(page_number=1, text=full_text)],
            total_pages=1,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            extra_metadata={
                "parser": "python-docx",
                "paragraph_count": str(len(document.paragraphs)),
                "table_count": str(len(document.tables)),
                "filename": filename,
            },
        )
